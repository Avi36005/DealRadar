"""Generates the 5 bundled sample documents in backend/sample_docs/.

Run once with:  python scripts/generate_sample_docs.py
Requires `reportlab` in addition to the runtime requirements (only needed
to regenerate these static sample files, not at app runtime):
    pip install reportlab
"""
from __future__ import annotations

from pathlib import Path

from docx import Document
from PIL import Image, ImageDraw, ImageFont
from reportlab.lib import colors
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import (
    Image as RLImage,
    PageBreak,
    Paragraph,
    SimpleDocTemplate,
    Spacer,
    Table,
    TableStyle,
)

OUT_DIR = Path(__file__).resolve().parent.parent / "sample_docs"
OUT_DIR.mkdir(parents=True, exist_ok=True)


def _font(size: int, bold: bool = False) -> ImageFont.ImageFont:
    candidates = (
        ["/System/Library/Fonts/Supplemental/Arial Bold.ttf"] if bold else []
    ) + [
        "/System/Library/Fonts/Supplemental/Arial.ttf",
        "/System/Library/Fonts/Helvetica.ttc",
    ]
    for candidate in candidates:
        try:
            return ImageFont.truetype(candidate, size)
        except OSError:
            continue
    return ImageFont.load_default()


# ─────────────────────────────────────────────────────────────────────────────
# 1. Scanned employee form — image-only PDF (no text layer, exercises OCR)
# ─────────────────────────────────────────────────────────────────────────────


def generate_scanned_form() -> None:
    width, height = 1240, 1754
    img = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(img)

    title_font = _font(34, bold=True)
    label_font = _font(22)
    body_font = _font(24)
    small_font = _font(18)

    margin = 90
    y = margin

    draw.text((margin, y), "EMPLOYEE INFORMATION FORM", fill="black", font=title_font)
    y += 60
    draw.text(
        (margin, y),
        "BFAI Technologies Pvt. Ltd.  |  Form No: HR-EIF-2024-0091",
        fill=(80, 80, 80),
        font=small_font,
    )
    y += 50
    draw.line((margin, y, width - margin, y), fill=(180, 180, 180), width=2)
    y += 40

    fields = [
        ("Full Name", "Ananya Sharma"),
        ("Employee ID", "EMP-10234"),
        ("Department", "Engineering"),
        ("Position", "Senior Software Engineer"),
        ("Date of Joining", "14 March 2022"),
        ("Reporting Manager", "Rohit Verma"),
        ("Work Location", "Bengaluru, India"),
        ("Employment Type", "Full-Time"),
        ("Date of Birth", "08 July 1994"),
        ("Personal Email", "ananya.sharma@example.com"),
    ]

    for label, value in fields:
        draw.text((margin, y), f"{label}:", fill="black", font=label_font)
        draw.line((margin + 320, y + 28, width - margin, y + 28), fill=(150, 150, 150), width=1)
        draw.text((margin + 340, y), value, fill=(20, 20, 20), font=body_font)
        y += 64

    y += 40
    draw.line((margin, y, width - margin, y), fill=(180, 180, 180), width=2)
    y += 40
    draw.text(
        (margin, y),
        "Employee Signature: ____________________     Date: ____________",
        fill="black",
        font=label_font,
    )
    y += 60
    draw.text((margin, y), "This document contains personal employee information.", fill=(80, 80, 80), font=small_font)
    y += 36
    draw.text((margin, y), "Classification: INTERNAL USE ONLY", fill=(150, 0, 0), font=label_font)

    out_path = OUT_DIR / "Employee Information Form (scanned).pdf"
    img.save(out_path, "PDF")
    print(f"Wrote {out_path}")


# ─────────────────────────────────────────────────────────────────────────────
# 2. Quarterly Sales Report — PDF with embedded tables
# ─────────────────────────────────────────────────────────────────────────────


def generate_sales_report() -> None:
    out_path = OUT_DIR / "Quarterly Sales Report.pdf"
    doc = SimpleDocTemplate(str(out_path), pagesize=letter, topMargin=0.75 * inch, bottomMargin=0.75 * inch)
    styles = getSampleStyleSheet()
    story = []

    story.append(Paragraph("BFAI Technologies — Quarterly Sales Report", styles["Title"]))
    story.append(Paragraph("Fiscal Year 2025 · Regional Performance Summary", styles["Normal"]))
    story.append(Spacer(1, 0.3 * inch))
    story.append(
        Paragraph(
            "This report summarizes regional sales performance for fiscal year 2025, covering "
            "revenue by quarter across five operating regions, along with a breakdown of the "
            "top performing products and customer growth trends. Figures are reported in USD thousands.",
            styles["BodyText"],
        )
    )
    story.append(Spacer(1, 0.25 * inch))

    story.append(Paragraph("Regional Revenue by Quarter (USD thousands)", styles["Heading2"]))
    region_data = [
        ["Region", "Q1", "Q2", "Q3", "Q4", "Total"],
        ["North America", "1,240", "1,310", "1,395", "1,480", "5,425"],
        ["Europe", "980", "1,015", "1,070", "1,125", "4,190"],
        ["Asia Pacific", "860", "905", "990", "1,110", "3,865"],
        ["Latin America", "320", "335", "360", "402", "1,417"],
        ["Middle East & Africa", "210", "228", "245", "268", "951"],
    ]
    t1 = Table(region_data, hAlign="LEFT")
    t1.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#16A34A")),
                ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
                ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
                ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#CCCCCC")),
                ("FONTSIZE", (0, 0), (-1, -1), 9),
                ("ALIGN", (1, 0), (-1, -1), "RIGHT"),
            ]
        )
    )
    story.append(t1)
    story.append(Spacer(1, 0.3 * inch))

    story.append(
        Paragraph(
            "North America remained the largest contributor at $5.43M for the year, growing "
            "steadily each quarter. Asia Pacific posted the strongest year-over-year growth at "
            "29%, driven by expansion in the enterprise segment.",
            styles["BodyText"],
        )
    )

    story.append(PageBreak())

    story.append(Paragraph("Top Performing Products", styles["Heading2"]))
    product_data = [
        ["Product", "Units Sold", "Revenue (USD)", "YoY Growth"],
        ["BFAI Parser Pro", "12,450", "$1,867,500", "+18%"],
        ["BFAI Classify Suite", "8,920", "$1,338,000", "+24%"],
        ["BFAI RAG Connector", "6,310", "$946,500", "+31%"],
        ["BFAI Voice Add-on", "3,105", "$310,500", "+45%"],
    ]
    t2 = Table(product_data, hAlign="LEFT")
    t2.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#2563EB")),
                ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
                ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
                ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#CCCCCC")),
                ("FONTSIZE", (0, 0), (-1, -1), 9),
                ("ALIGN", (1, 0), (-1, -1), "RIGHT"),
            ]
        )
    )
    story.append(t2)
    story.append(Spacer(1, 0.3 * inch))

    story.append(Paragraph("Customer Growth", styles["Heading2"]))
    cust_data = [
        ["Quarter", "New Customers", "Churned Customers", "Net Growth"],
        ["Q1 2025", "145", "32", "+113"],
        ["Q2 2025", "168", "28", "+140"],
        ["Q3 2025", "182", "35", "+147"],
        ["Q4 2025", "201", "30", "+171"],
    ]
    t3 = Table(cust_data, hAlign="LEFT")
    t3.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#D97706")),
                ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
                ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
                ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#CCCCCC")),
                ("FONTSIZE", (0, 0), (-1, -1), 9),
                ("ALIGN", (1, 0), (-1, -1), "RIGHT"),
            ]
        )
    )
    story.append(t3)
    story.append(Spacer(1, 0.25 * inch))
    story.append(
        Paragraph(
            "Net new customer growth accelerated through the year, with Q4 2025 adding 171 net "
            "new accounts — the strongest quarter on record. This report is classified as "
            "Confidential and intended for internal financial planning use only.",
            styles["BodyText"],
        )
    )

    doc.build(story)
    print(f"Wrote {out_path}")


# ─────────────────────────────────────────────────────────────────────────────
# 3. Network Architecture Overview — image-heavy PDF (text + embedded diagram)
# ─────────────────────────────────────────────────────────────────────────────


def _make_network_diagram(path: Path) -> None:
    width, height = 1000, 600
    img = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(img)
    font = _font(20)
    bold = _font(22, bold=True)

    boxes = {
        "Internet": (40, 40, 240, 110),
        "Edge Firewall": (400, 40, 620, 110),
        "Load Balancer": (760, 40, 960, 110),
        "Core Switch": (400, 220, 620, 290),
        "App Server 1": (160, 400, 380, 470),
        "App Server 2": (420, 400, 640, 470),
        "Database Cluster": (680, 400, 940, 470),
    }

    for label, (x0, y0, x1, y1) in boxes.items():
        draw.rectangle((x0, y0, x1, y1), outline="black", width=3)
        draw.text((x0 + 12, y0 + (y1 - y0) // 2 - 12), label, fill="black", font=bold)

    def center(box: str) -> tuple[int, int]:
        x0, y0, x1, y1 = boxes[box]
        return (x0 + x1) // 2, (y0 + y1) // 2

    edges = [
        ("Internet", "Edge Firewall"),
        ("Edge Firewall", "Load Balancer"),
        ("Edge Firewall", "Core Switch"),
        ("Core Switch", "App Server 1"),
        ("Core Switch", "App Server 2"),
        ("Core Switch", "Database Cluster"),
    ]
    for a, b in edges:
        draw.line((*center(a), *center(b)), fill=(100, 100, 100), width=2)

    draw.text((40, 540), "Fig. 1 - BFAI Production Network Topology (Region: ap-south-1)", fill=(80, 80, 80), font=font)
    img.save(path, "PNG")


def generate_network_overview() -> None:
    diagram_path = OUT_DIR / "_network_diagram_tmp.png"
    _make_network_diagram(diagram_path)

    out_path = OUT_DIR / "Network Architecture Overview.pdf"
    doc = SimpleDocTemplate(str(out_path), pagesize=letter, topMargin=0.75 * inch, bottomMargin=0.75 * inch)
    styles = getSampleStyleSheet()
    story = []

    story.append(Paragraph("Network Architecture Overview", styles["Title"]))
    story.append(Spacer(1, 0.15 * inch))
    story.append(
        Paragraph(
            "This document describes the production network architecture for the BFAI platform, "
            "including perimeter security, load balancing, and internal service segmentation. "
            "All production traffic enters through the edge firewall before reaching the load "
            "balancer tier.",
            styles["BodyText"],
        )
    )
    story.append(Spacer(1, 0.2 * inch))
    story.append(RLImage(str(diagram_path), width=6.5 * inch, height=3.9 * inch))
    story.append(Spacer(1, 0.2 * inch))
    story.append(
        Paragraph(
            "The core switch routes internal traffic between the application server tier and the "
            "database cluster. Application servers are deployed across two availability zones for "
            "redundancy, and the database cluster replicates synchronously between zones.",
            styles["BodyText"],
        )
    )

    story.append(PageBreak())
    story.append(Paragraph("Security Zones", styles["Heading2"]))
    story.append(
        Paragraph(
            "The network is divided into three security zones: the public zone (internet-facing "
            "load balancers), the application zone (app servers, internal APIs), and the data zone "
            "(database cluster, object storage). Traffic between zones is restricted by firewall "
            "rules that allow only the minimum required ports.",
            styles["BodyText"],
        )
    )
    story.append(Spacer(1, 0.15 * inch))
    story.append(
        Paragraph(
            "Access to the data zone is limited to application servers and the operations team via "
            "a bastion host. All administrative access requires multi-factor authentication and is "
            "logged for audit purposes.",
            styles["BodyText"],
        )
    )

    doc.build(story)
    diagram_path.unlink(missing_ok=True)
    print(f"Wrote {out_path}")


# ─────────────────────────────────────────────────────────────────────────────
# 4. Remote Work Policy — plain text
# ─────────────────────────────────────────────────────────────────────────────


def generate_remote_work_policy() -> None:
    content = """REMOTE WORK POLICY
BFAI Technologies Pvt. Ltd.
Effective Date: January 1, 2026
Document Classification: Public

1. PURPOSE
This policy establishes guidelines for employees who work remotely on a full-time or
hybrid basis. It is intended to ensure consistency, productivity, and security across
all remote work arrangements while supporting employee flexibility and wellbeing.

2. ELIGIBILITY
Remote work arrangements are available to employees whose role does not require
regular physical presence at a company facility, subject to manager approval. New
employees typically complete their first 90 days on-site or in a hybrid arrangement
before transitioning to a fully remote schedule, at the discretion of their manager.

3. WORK HOURS AND AVAILABILITY
Remote employees are expected to maintain the same core working hours as their team,
generally between 10:00 AM and 4:00 PM in their local time zone, to support
collaboration across distributed teams. Employees should keep their calendar status
up to date and respond to messages within a reasonable timeframe during work hours.

4. EQUIPMENT AND WORKSPACE
The company provides a laptop, monitor, and necessary peripherals for remote
employees. Employees are responsible for maintaining a dedicated, ergonomic
workspace with a stable internet connection of at least 25 Mbps download speed.
Equipment remains company property and must be returned upon termination of
employment.

5. COMMUNICATION EXPECTATIONS
Remote employees must attend scheduled team meetings via video call with cameras on,
unless otherwise agreed with their manager. Daily standups, sprint planning, and
retrospectives are considered mandatory unless the employee is on approved leave.

6. SECURITY REQUIREMENTS
All remote work must be conducted on company-managed devices with full-disk
encryption, automatic screen lock after 5 minutes of inactivity, and up-to-date
antivirus software. Employees must connect to internal systems via the company VPN
and must not use public, unsecured Wi-Fi networks to access confidential systems.
Sharing login credentials or company devices with family members or others is
strictly prohibited.

7. PERFORMANCE EXPECTATIONS
Remote employees are evaluated on the same performance criteria as on-site
employees, based on deliverables, quality of work, and collaboration, rather than
hours logged. Managers will conduct regular one-on-one check-ins to discuss
workload, priorities, and any support needed.

8. EXPENSES AND REIMBURSEMENT
Employees may submit monthly reimbursement requests of up to $50 for internet
connectivity costs, subject to finance team approval. Additional equipment
requests beyond the standard kit require manager and IT approval.

9. POLICY REVIEW
This policy will be reviewed annually by the People Operations team and may be
updated to reflect changes in business needs, labor regulations, or company
strategy. Employees will be notified of material changes at least 30 days before
they take effect.

For questions about this policy, contact people-ops@bfai.example.com.
"""
    out_path = OUT_DIR / "Remote Work Policy.txt"
    out_path.write_text(content, encoding="utf-8")
    print(f"Wrote {out_path}")


# ─────────────────────────────────────────────────────────────────────────────
# 5. Project Proposal — multi-page DOCX with a table
# ─────────────────────────────────────────────────────────────────────────────


def generate_project_proposal() -> None:
    doc = Document()
    doc.add_heading("Project Proposal: AI-Powered Document Intelligence Platform", level=0)

    doc.add_heading("1. Executive Summary", level=1)
    doc.add_paragraph(
        "This proposal outlines the development of an AI-powered document intelligence "
        "platform that automatically parses, classifies, and enables natural-language "
        "question answering over heterogeneous document collections, including scanned "
        "PDFs, handwritten forms, and structured reports. The platform will reduce manual "
        "document review time by an estimated 60% across pilot teams."
    )

    doc.add_heading("2. Background", level=1)
    doc.add_paragraph(
        "Internal teams currently spend significant time manually searching through "
        "contracts, invoices, reports, and policy documents stored across shared drives. "
        "Existing keyword search tools do not understand document structure, cannot read "
        "scanned content, and provide no way to verify where an answer came from."
    )

    doc.add_heading("3. Objectives", level=1)
    for item in [
        "Automatically extract text, tables, and page images from any uploaded document.",
        "Classify each document by type, topic, sensitivity, and content characteristics.",
        "Provide a conversational interface that answers questions with page-level citations.",
        "Support bulk ingestion with real-time per-document processing status.",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("4. Scope", level=1)
    doc.add_paragraph(
        "The initial release covers PDF (including scanned and image-based), DOCX, and "
        "plain text formats. Out of scope for this phase: spreadsheets, presentation "
        "files, and email ingestion, which are planned for a future release."
    )

    doc.add_heading("5. Timeline", level=1)
    table = doc.add_table(rows=1, cols=3)
    table.style = "Light Grid Accent 1"
    hdr = table.rows[0].cells
    hdr[0].text, hdr[1].text, hdr[2].text = "Phase", "Duration", "Deliverable"
    for phase, duration, deliverable in [
        ("Discovery & Design", "2 weeks", "Architecture document, schema definitions"),
        ("Core Parsing & Classification", "4 weeks", "Parser service, classifier service"),
        ("Agentic RAG & Chat UI", "3 weeks", "Retrieval pipeline, chat interface with citations"),
        ("Bulk Upload & Hardening", "2 weeks", "Bulk upload UI, security review, deployment"),
    ]:
        cells = table.add_row().cells
        cells[0].text, cells[1].text, cells[2].text = phase, duration, deliverable

    doc.add_heading("6. Budget", level=1)
    table2 = doc.add_table(rows=1, cols=2)
    table2.style = "Light Grid Accent 1"
    hdr2 = table2.rows[0].cells
    hdr2[0].text, hdr2[1].text = "Item", "Estimated Cost (USD)"
    for item, cost in [
        ("Engineering (3 FTE x 11 weeks)", "$66,000"),
        ("Cloud infrastructure (pilot)", "$1,200"),
        ("Third-party API usage (free tier)", "$0"),
        ("Contingency (10%)", "$6,720"),
    ]:
        cells = table2.add_row().cells
        cells[0].text, cells[1].text = item, cost

    doc.add_heading("7. Risks", level=1)
    for item in [
        "OCR accuracy may vary for low-quality scans or handwriting.",
        "Free-tier LLM rate limits may require request batching under heavy load.",
        "Sensitive documents require careful handling of storage permissions and access control.",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("8. Conclusion", level=1)
    doc.add_paragraph(
        "This platform addresses a clear operational need with a modest budget and a phased "
        "delivery plan. Approval is requested to proceed with the Discovery & Design phase "
        "by the start of next quarter."
    )

    out_path = OUT_DIR / "Project Proposal.docx"
    doc.save(out_path)
    print(f"Wrote {out_path}")


if __name__ == "__main__":
    generate_scanned_form()
    generate_sales_report()
    generate_network_overview()
    generate_remote_work_policy()
    generate_project_proposal()
