"""Document parsing: extracts text, tables, and rendered page images.

Supported inputs:
  - PDF (text, scanned/image-only, handwritten, tables, image-heavy)
  - DOCX (python-docx)
  - TXT (plain text)

Every page produces a `PageData` record with `extracted_text`,
`page_image_path` (PNG rendered to disk), and any `tables` found
(stored as structured headers/rows JSON, never flattened to text).
"""
from __future__ import annotations

import json
import logging
import textwrap
from pathlib import Path
from typing import List

import pdfplumber
from PIL import Image, ImageDraw, ImageFont
from docx import Document as DocxDocument

from config import PAGE_IMAGES_DIR, PARSED_DIR
from models.schemas import PageData, ParsedDocument, TableData

logger = logging.getLogger(__name__)

OCR_TEXT_THRESHOLD = 30  # below this many chars, treat page as scanned/image-only
PAGE_RENDER_DPI = 150
TEXT_PAGE_CHARS = 1200   # ~chars per rendered "page" for txt/docx


# ─────────────────────────────────────────────────────────────────────────────
# Public entry point
# ─────────────────────────────────────────────────────────────────────────────


def parse_document(document_id: str, file_path: Path, document_name: str) -> ParsedDocument:
    """Parse `file_path` and persist the result to storage/parsed/{document_id}.json."""
    ext = file_path.suffix.lower()
    image_dir = PAGE_IMAGES_DIR / document_id
    image_dir.mkdir(parents=True, exist_ok=True)

    if ext == ".pdf":
        pages = _parse_pdf(document_id, file_path, image_dir)
    elif ext == ".docx":
        pages = _parse_docx(document_id, file_path, image_dir)
    elif ext == ".txt":
        pages = _parse_txt(document_id, file_path, image_dir)
    else:
        raise ValueError(f"Unsupported file type: {ext}")

    if not pages:
        # Always guarantee at least one page so downstream steps have something
        pages = [_blank_page(document_id, 1, image_dir)]

    parsed = ParsedDocument(document_id=document_id, document_name=document_name, pages=pages)
    _persist(parsed)
    return parsed


def load_parsed(document_id: str) -> ParsedDocument | None:
    path = PARSED_DIR / f"{document_id}.json"
    if not path.exists():
        return None
    with open(path, "r", encoding="utf-8") as f:
        return ParsedDocument.model_validate(json.load(f))


# ─────────────────────────────────────────────────────────────────────────────
# PDF
# ─────────────────────────────────────────────────────────────────────────────


def _parse_pdf(document_id: str, file_path: Path, image_dir: Path) -> List[PageData]:
    pages: List[PageData] = []

    rendered_images = _render_pdf_pages(file_path, image_dir)

    with pdfplumber.open(file_path) as pdf:
        for i, pdf_page in enumerate(pdf.pages):
            page_number = i + 1
            image_path = rendered_images.get(page_number)

            text = (pdf_page.extract_text() or "").strip()

            tables = _extract_tables(pdf_page)

            # Scanned / handwritten pages have little or no selectable text —
            # OCR the rendered page image. Image-heavy pages may already have
            # body text but also embedded diagrams/photos with text inside
            # them, so OCR those too and append anything new.
            is_scanned = len(text) < OCR_TEXT_THRESHOLD
            has_embedded_images = bool(getattr(pdf_page, "images", None))
            if (is_scanned or has_embedded_images) and image_path:
                ocr_text = _ocr_image(image_dir / image_path)
                if ocr_text and ocr_text not in text:
                    text = (text + "\n" + ocr_text).strip() if text else ocr_text

            pages.append(
                PageData(
                    document_id=document_id,
                    page_number=page_number,
                    extracted_text=text,
                    page_image_path=_rel_image_path(document_id, image_path),
                    tables=tables,
                    is_scanned=is_scanned,
                    has_images=has_embedded_images,
                )
            )

    return pages


def _render_pdf_pages(file_path: Path, image_dir: Path) -> dict[int, str]:
    """Render every PDF page to a PNG. Returns {page_number: filename}."""
    try:
        from pdf2image import convert_from_path

        images = convert_from_path(str(file_path), dpi=PAGE_RENDER_DPI)
    except Exception as exc:  # poppler missing, corrupt PDF, etc.
        logger.warning("pdf2image failed for %s: %s", file_path, exc)
        return {}

    result: dict[int, str] = {}
    for i, image in enumerate(images):
        page_number = i + 1
        filename = f"page_{page_number:03d}.png"
        image.convert("RGB").save(image_dir / filename, "PNG")
        result[page_number] = filename
    return result


def _extract_tables(pdf_page) -> List[TableData]:
    tables: List[TableData] = []
    try:
        raw_tables = pdf_page.extract_tables()
    except Exception as exc:
        logger.warning("Table extraction failed: %s", exc)
        return tables

    for raw in raw_tables:
        if not raw:
            continue
        clean_rows = [["" if cell is None else str(cell).strip() for cell in row] for row in raw]
        headers, *rows = clean_rows
        tables.append(TableData(headers=headers, rows=rows))
    return tables


def _ocr_image(image_path: Path) -> str:
    try:
        import pytesseract

        return pytesseract.image_to_string(Image.open(image_path)).strip()
    except Exception as exc:
        logger.warning("OCR failed for %s: %s", image_path, exc)
        return ""


# ─────────────────────────────────────────────────────────────────────────────
# DOCX
# ─────────────────────────────────────────────────────────────────────────────


def _parse_docx(document_id: str, file_path: Path, image_dir: Path) -> List[PageData]:
    doc = DocxDocument(str(file_path))

    full_text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
    text_pages = _chunk_text(full_text)

    tables: List[TableData] = []
    for table in doc.tables:
        rows = [[cell.text.strip() for cell in row.cells] for row in table.rows]
        if not rows:
            continue
        headers, *body = rows
        tables.append(TableData(headers=headers, rows=body))

    if not text_pages:
        text_pages = [""]

    pages: List[PageData] = []
    for i, page_text in enumerate(text_pages):
        page_number = i + 1
        filename = f"page_{page_number:03d}.png"
        _render_text_image(page_text, image_dir / filename, title=file_path.name)
        pages.append(
            PageData(
                document_id=document_id,
                page_number=page_number,
                extracted_text=page_text,
                page_image_path=_rel_image_path(document_id, filename),
                tables=tables if i == 0 else [],
            )
        )
    return pages


# ─────────────────────────────────────────────────────────────────────────────
# TXT
# ─────────────────────────────────────────────────────────────────────────────


def _parse_txt(document_id: str, file_path: Path, image_dir: Path) -> List[PageData]:
    raw = file_path.read_bytes()
    text = raw.decode("utf-8", errors="ignore")
    text_pages = _chunk_text(text) or [""]

    pages: List[PageData] = []
    for i, page_text in enumerate(text_pages):
        page_number = i + 1
        filename = f"page_{page_number:03d}.png"
        _render_text_image(page_text, image_dir / filename, title=file_path.name)
        pages.append(
            PageData(
                document_id=document_id,
                page_number=page_number,
                extracted_text=page_text,
                page_image_path=_rel_image_path(document_id, filename),
                tables=[],
            )
        )
    return pages


# ─────────────────────────────────────────────────────────────────────────────
# Shared helpers
# ─────────────────────────────────────────────────────────────────────────────


def _chunk_text(text: str, max_chars: int = TEXT_PAGE_CHARS) -> List[str]:
    """Split text into page-sized chunks, breaking on paragraph/line boundaries."""
    text = text.strip()
    if not text:
        return []

    paragraphs = text.split("\n")
    pages: List[str] = []
    current = ""

    for para in paragraphs:
        candidate = f"{current}\n{para}" if current else para
        if len(candidate) > max_chars and current:
            pages.append(current)
            current = para
        else:
            current = candidate

        # A single huge paragraph still needs splitting.
        while len(current) > max_chars:
            pages.append(current[:max_chars])
            current = current[max_chars:]

    if current:
        pages.append(current)

    return pages


def _render_text_image(text: str, out_path: Path, title: str = "") -> None:
    """Render a block of text onto a page-sized PNG for use as a citation thumbnail."""
    width, height = 1240, 1754  # ~A4 at 150dpi
    margin = 60
    image = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(image)

    font = _load_font(18)
    title_font = _load_font(22)

    y = margin
    if title:
        draw.text((margin, y), title, fill="black", font=title_font)
        y += 36
        draw.line((margin, y, width - margin, y), fill=(220, 220, 220), width=2)
        y += 20

    wrap_width = 110
    for line in text.split("\n"):
        wrapped = textwrap.wrap(line, width=wrap_width) or [""]
        for wrapped_line in wrapped:
            if y > height - margin:
                break
            draw.text((margin, y), wrapped_line, fill="black", font=font)
            y += 26

    image.save(out_path, "PNG")


def _load_font(size: int) -> ImageFont.ImageFont:
    for candidate in (
        "/System/Library/Fonts/Supplemental/Arial.ttf",
        "/System/Library/Fonts/Helvetica.ttc",
        "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf",
    ):
        try:
            return ImageFont.truetype(candidate, size)
        except OSError:
            continue
    return ImageFont.load_default()


def _blank_page(document_id: str, page_number: int, image_dir: Path) -> PageData:
    filename = f"page_{page_number:03d}.png"
    _render_text_image("", image_dir / filename)
    return PageData(
        document_id=document_id,
        page_number=page_number,
        extracted_text="",
        page_image_path=_rel_image_path(document_id, filename),
        tables=[],
    )


def _rel_image_path(document_id: str, filename: str | None) -> str:
    if not filename:
        filename = "page_001.png"
    return f"storage/page_images/{document_id}/{filename}"


def _persist(parsed: ParsedDocument) -> None:
    PARSED_DIR.mkdir(parents=True, exist_ok=True)
    out_path = PARSED_DIR / f"{parsed.document_id}.json"
    with open(out_path, "w", encoding="utf-8") as f:
        json.dump(parsed.model_dump(), f, indent=2)
